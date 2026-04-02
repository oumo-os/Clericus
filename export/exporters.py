"""
export/exporters.py
-------------------
Export the finished document tree to txt / md / html / pdf / docx.

Key fixes:
  - All renderers now recurse on node["children"] (not node["body"] as list),
    which is what recursive_drafter produces.
  - Body text is taken from node["body_text"] with fallback to node["body"].
  - dedupe_references + safe stringify prevent dict-ref crashes in docx.
  - export_to_html / export_to_pdf fall back gracefully when the
    templates/ directory doesn't exist.
"""

import os
from pathlib import Path
from typing import Any, Dict, List

from utils.references import format_references_md, dedupe_references
from utils.text_tools import normalize_whitespace


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _body(node: Dict[str, Any]) -> str:
    """Return the best available body text string from a node."""
    return normalize_whitespace(
        (node.get("body_text") or node.get("body") or "").strip()
    )


def _ref_str(ref: Any) -> str:
    """Safely convert a reference (string or dict) to a display string."""
    if isinstance(ref, dict):
        author   = ref.get("author", ref.get("source", "Unknown"))
        year     = ref.get("year", "n.d.")
        title    = ref.get("title", "")
        filename = ref.get("filename", "")
        page     = ref.get("page", "")
        parts = [f"{author} ({year})"]
        if title:
            parts.append(title)
        if page and page != "?":
            parts.append(f"p.{page}")
        if filename:
            parts.append(f"[{filename}]")
        return " — ".join(parts)
    return str(ref)


def _collect_references(node: Dict[str, Any]) -> List[Any]:
    """Recursively collect all references from a document tree."""
    refs = list(node.get("references") or [])
    for child in node.get("children") or []:
        refs.extend(_collect_references(child))
    return dedupe_references(refs)


# ---------------------------------------------------------------------------
# Plain text
# ---------------------------------------------------------------------------

def export_to_txt(document: Dict[str, Any], output_path: str) -> None:
    lines: List[str] = []

    def _walk(node: Dict[str, Any], depth: int = 0) -> None:
        indent = "  " * depth
        lines.append(f"{indent}{node.get('title', '')}")
        opening = normalize_whitespace((node.get("openning") or "").strip())
        if opening:
            lines.append(f"{indent}{opening}\n")
        body = _body(node)
        if body:
            lines.append(f"{indent}{body}\n")
        for child in node.get("children") or []:
            _walk(child, depth + 1)
        closing = normalize_whitespace((node.get("closing") or "").strip())
        if closing:
            lines.append(f"{indent}{closing}\n")

    _walk(document)
    lines.append("\nReferences:")
    for ref in _collect_references(document):
        lines.append(f"  - {_ref_str(ref)}")

    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    Path(output_path).write_text("\n".join(lines), encoding="utf-8")


# ---------------------------------------------------------------------------
# Markdown
# ---------------------------------------------------------------------------

def export_to_md(document: Dict[str, Any], output_path: str) -> None:
    lines: List[str] = []

    def _walk(node: Dict[str, Any], depth: int = 1) -> None:
        prefix = "#" * min(depth, 6)
        lines.append(f"{prefix} {node.get('title', '')}\n")
        opening = (node.get("openning") or "").strip()
        if opening:
            lines.append(f"{opening}\n")
        body = _body(node)
        if body:
            lines.append(f"{body}\n")
        for child in node.get("children") or []:
            _walk(child, depth + 1)
        closing = (node.get("closing") or "").strip()
        if closing:
            lines.append(f"*{closing}*\n")

    _walk(document)
    lines.append("\n## References\n")
    for ref in _collect_references(document):
        lines.append(f"- {_ref_str(ref)}")

    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    Path(output_path).write_text("\n".join(lines), encoding="utf-8")


# ---------------------------------------------------------------------------
# HTML
# ---------------------------------------------------------------------------

def export_to_html(document: Dict[str, Any], output_path: str) -> None:
    """
    Renders via Jinja2 template if templates/document.html exists,
    otherwise falls back to a minimal inline HTML render.
    """
    TEMPLATES_DIR = Path(__file__).parent / "templates"

    if (TEMPLATES_DIR / "document.html").exists():
        from jinja2 import Environment, FileSystemLoader
        env = Environment(loader=FileSystemLoader(str(TEMPLATES_DIR)))
        html = env.get_template("document.html").render(document=document)
    else:
        html = _render_html_inline(document)

    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    Path(output_path).write_text(html, encoding="utf-8")


def _render_html_inline(document: Dict[str, Any]) -> str:
    parts = ["<!DOCTYPE html><html><head><meta charset='utf-8'>",
             f"<title>{document.get('title','Document')}</title></head><body>"]

    def _walk(node: Dict[str, Any], depth: int = 1) -> None:
        tag = f"h{min(depth, 6)}"
        parts.append(f"<{tag}>{node.get('title','')}</{tag}>")
        opening = (node.get("openning") or "").strip()
        if opening:
            parts.append(f"<p>{opening}</p>")
        body = _body(node)
        if body:
            for para in body.split("\n\n"):
                if para.strip():
                    parts.append(f"<p>{para.strip()}</p>")
        for child in node.get("children") or []:
            _walk(child, depth + 1)
        closing = (node.get("closing") or "").strip()
        if closing:
            parts.append(f"<p><em>{closing}</em></p>")

    _walk(document)
    refs = _collect_references(document)
    if refs:
        parts.append("<h2>References</h2><ul>")
        for ref in refs:
            parts.append(f"<li>{_ref_str(ref)}</li>")
        parts.append("</ul>")
    parts.append("</body></html>")
    return "\n".join(parts)


# ---------------------------------------------------------------------------
# PDF  (via WeasyPrint → HTML)
# ---------------------------------------------------------------------------

def export_to_pdf(document: Dict[str, Any], output_path: str) -> None:
    # Build HTML first, then convert
    html_path = output_path.replace(".pdf", "_tmp.html")
    export_to_html(document, html_path)
    try:
        from weasyprint import HTML as WP
        WP(filename=html_path).write_pdf(output_path)
    finally:
        try:
            os.unlink(html_path)
        except OSError:
            pass


# ---------------------------------------------------------------------------
# DOCX
# ---------------------------------------------------------------------------

def export_to_docx(document: Dict[str, Any], output_path: str) -> None:
    from docx import Document as DocxDocument
    from docx.shared import Pt

    doc = DocxDocument()
    doc.add_heading(document.get("title", "Document"), level=0)

    opening = (document.get("openning") or "").strip()
    if opening:
        doc.add_paragraph(normalize_whitespace(opening))

    def _walk(node: Dict[str, Any], level: int = 1) -> None:
        doc.add_heading(node.get("title", ""), level=min(level, 9))
        opening = normalize_whitespace((node.get("openning") or "").strip())
        if opening:
            doc.add_paragraph(opening)
        body = _body(node)
        if body:
            for para in body.split("\n\n"):
                if para.strip():
                    doc.add_paragraph(para.strip())
        for child in node.get("children") or []:
            _walk(child, level + 1)
        closing = normalize_whitespace((node.get("closing") or "").strip())
        if closing:
            p = doc.add_paragraph(closing)
            p.runs[0].italic = True if p.runs else None

    for section in document.get("children") or []:
        _walk(section)

    # Top-level body_text fallback (single-node documents)
    root_body = _body(document)
    if root_body and not document.get("children"):
        doc.add_paragraph(root_body)

    # References appendix
    refs = _collect_references(document)
    if refs:
        doc.add_heading("References", level=1)
        for ref in refs:
            p = doc.add_paragraph(_ref_str(ref), style="List Bullet")
            for run in p.runs:
                run.font.size = Pt(10)

    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)


# ---------------------------------------------------------------------------
# Dispatcher
# ---------------------------------------------------------------------------

def export_document(document: Dict[str, Any], fmt: str, output_path: str) -> None:
    """Route to the correct exporter."""
    fmt = fmt.lower().strip()
    dispatch = {
        "txt":  export_to_txt,
        "md":   export_to_md,
        "html": export_to_html,
        "pdf":  export_to_pdf,
        "docx": export_to_docx,
    }
    fn = dispatch.get(fmt)
    if fn is None:
        raise ValueError(f"Unsupported export format: '{fmt}'. Choose from {list(dispatch)}")
    fn(document, output_path)
